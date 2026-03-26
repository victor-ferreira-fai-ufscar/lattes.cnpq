import os
from docx import Document
from docx.shared import Pt
import datetime

def create_lattes_docx(nome: str, dados_gemini: dict, output_dir: str) -> str:
    """
    Gera um relatório DOCX estruturado a partir do JSON gerado pelo Gemini e salva no diretório de output.
    Retorna o caminho do arquivo gerado.
    """
    os.makedirs(output_dir, exist_ok=True)
    filename = f"{nome.replace(' ', '_')}_Lattes_{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    filepath = os.path.join(output_dir, filename)
    
    doc = Document()
    
    # Título Principal
    titulo = doc.add_heading(f"Relatório de Perfil Lattes: {nome}", 0)
    titulo.alignment = 1 # Center
    
    # Resumo Executivo
    doc.add_heading("Resumo Executivo", level=1)
    p_resumo = doc.add_paragraph(dados_gemini.get('resumo', 'Não foi possível gerar o resumo.'))
    p_resumo.style.font.size = Pt(11)
    
    # Vínculo Institucional
    doc.add_heading("Vínculo Institucional", level=1)
    vinculo = dados_gemini.get('vinculo_institucional', '')
    doc.add_paragraph(vinculo if vinculo else "Nenhum vínculo detectado.")
    
    # Trajetória Acadêmica
    doc.add_heading("Trajetória Acadêmica", level=1)
    
    trajetoria = [
        ("Graduação", dados_gemini.get('graduacao', '')),
        ("Mestrado", dados_gemini.get('mestrado', '')),
        ("Doutorado", dados_gemini.get('doutorado', '')),
        ("Pós-Doutorado", dados_gemini.get('pos_doutorado', ''))
    ]
    
    for nivel, texto in trajetoria:
        if texto:
            p = doc.add_paragraph()
            p.add_run(f"{nivel}: ").bold = True
            p.add_run(texto)
            
    doc.add_page_break()
    
    # Rodapé simples
    section = doc.sections[-1]
    footer = section.footer
    footer.paragraphs[0].text = f"Gerado automaticamente pelo Lattes Scraper Bot em {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}"

    doc.save(filepath)
    return filepath
