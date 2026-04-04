import csv
import json
import os
import re
import zipfile
from base64 import b64encode
from dataclasses import asdict, dataclass
from datetime import date, datetime
from html import escape
from io import BytesIO, StringIO
from typing import Literal

from .scraper import _extrair_texto_pdf_bytes
from .storage import (
    download_storage_file_bytes,
    list_storage_files,
    upload_file_bytes,
)
from ..libs.filename import slugify_nome

OutputFormat = Literal["pdf", "docx", "json", "html", "csv", "all"]

DEFAULT_OUTPUT_FORMAT: OutputFormat = "docx"
_FORMAT_ORDER: tuple[OutputFormat, ...] = ("docx", "json", "html", "csv", "pdf")
_SUPPORTED_FORMATS: set[str] = set(_FORMAT_ORDER) | {"all"}
_MANIFEST_FILENAME = "manifest.json"
_SUMMARY_TEMPLATE_NAME = "vitrine-resumo-v2-com-foto (refs: claudia_martinez.docx, Antonio José Gonçalves da Cruz devolutiva 1.docx)"
_DEFAULT_STORAGE_ROOT = "structured/outputs"
_DEFAULT_TEMPLATE_VERSION = "v3"


@dataclass(frozen=True)
class GeneratedArtifact:
    format: str
    filename: str
    relative_path: str
    download_url: str
    content_type: str


@dataclass(frozen=True)
class GeneratedArtifactBundle:
    output_format: str
    output_directory: str
    output_label: str
    generated_files: list[GeneratedArtifact]
    extracted_text_length: int
    template_name: str | None = None
    zip_file: GeneratedArtifact | None = None
    artifacts_cache_status: str = "miss"


@dataclass(frozen=True)
class ExportSummaryContext:
    summary_paragraph: str
    keywords: str
    knowledge_area: str
    social_links: list[str]
    focus_topics: str
    ods: list[str]
    source_excerpt: str
    has_profile_photo: bool = False


def normalize_output_format(value: str | None) -> OutputFormat:
    normalized = (value or DEFAULT_OUTPUT_FORMAT).strip().lower()
    if normalized not in _SUPPORTED_FORMATS:
        supported = ", ".join(sorted(_SUPPORTED_FORMATS))
        raise ValueError(
            f"Formato de saída inválido: '{value}'. Use um destes: {supported}."
        )
    return normalized  # type: ignore[return-value]


def expand_output_formats(output_format: OutputFormat) -> list[OutputFormat]:
    if output_format == "all":
        return list(_FORMAT_ORDER)
    return [output_format]


def _artifact_storage_root() -> str:
    return (
        os.getenv("SUPABASE_STORAGE_STRUCTURED_FOLDER", _DEFAULT_STORAGE_ROOT).strip()
        or _DEFAULT_STORAGE_ROOT
    ).strip("/")


def _artifact_template_version() -> str:
    return (
        os.getenv("LATTES_EXPORT_TEMPLATE_VERSION", _DEFAULT_TEMPLATE_VERSION).strip()
        or _DEFAULT_TEMPLATE_VERSION
    )


def _normalize_display_name(nome: str) -> str:
    sanitized = " ".join(nome.split())
    sanitized = sanitized.replace("/", "-").replace("\\", "-")
    return sanitized or "Docente"


def build_curriculo_output_label(nome: str, ultima_atualizacao: date) -> str:
    return f"{_normalize_display_name(nome)} - {ultima_atualizacao.isoformat()}"


def build_curriculo_docx_filename(nome: str, ultima_atualizacao: date) -> str:
    normalized = re.sub(
        r"-+", "-", _normalize_display_name(nome).replace(" ", "-")
    ).strip("-")
    return f"perfil-vitrine-{normalized}-{ultima_atualizacao.isoformat()}.docx"


def build_curriculo_storage_folder(nome: str, ultima_atualizacao: date) -> str:
    return "/".join(
        [
            _artifact_storage_root(),
            _artifact_template_version(),
            "curriculos",
            slugify_nome(nome),
            ultima_atualizacao.isoformat(),
        ]
    )


def build_batch_storage_folder(batch_id: str) -> str:
    return "/".join(
        [_artifact_storage_root(), _artifact_template_version(), "lotes", batch_id]
    )


def create_batch_storage_target() -> tuple[str, str]:
    batch_id = f"lote-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
    return build_batch_storage_folder(batch_id), batch_id


def _package_zip_filename(nome: str, ultima_atualizacao: date) -> str:
    return f"pacote-{slugify_nome(nome)}-{ultima_atualizacao.isoformat()}.zip"


def _artifact_file_specs(
    nome: str, ultima_atualizacao: date
) -> dict[str, tuple[str, str]]:
    return {
        "pdf": ("curriculo-lattes.pdf", "application/pdf"),
        "docx": (
            build_curriculo_docx_filename(nome, ultima_atualizacao),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
        "json": ("dados-extraidos.json", "application/json"),
        "html": ("curriculo-lattes.html", "text/html"),
        "csv": ("dados-extraidos.csv", "text/csv"),
    }


def _clean_text_lines(texto: str) -> list[str]:
    return [
        line.strip() for line in texto.replace("\r", "").splitlines() if line.strip()
    ]


def _extract_urls(lines: list[str]) -> list[str]:
    url_pattern = re.compile(r"https?://\S+", re.IGNORECASE)
    found: list[str] = []
    seen: set[str] = set()

    for line in lines:
        for match in url_pattern.findall(line):
            normalized = match.rstrip(".,);]")
            if normalized.lower() in seen:
                continue
            seen.add(normalized.lower())
            found.append(normalized)

    return found[:8]


def _extract_tagged_block(lines: list[str], tag_candidates: list[str]) -> str | None:
    normalized_tags = tuple(candidate.lower() for candidate in tag_candidates)
    for index, line in enumerate(lines):
        normalized = line.lower()
        if any(tag in normalized for tag in normalized_tags):
            if ":" in line:
                tail = line.split(":", 1)[1].strip()
                if tail:
                    return tail

            collected: list[str] = []
            for next_line in lines[index + 1 : index + 5]:
                next_normalized = next_line.lower()
                if next_normalized.startswith("campo ") or next_normalized.endswith(
                    ":"
                ):
                    break
                if next_line.isupper() and len(next_line.split()) <= 6:
                    break
                collected.append(next_line)
            if collected:
                return " ".join(collected).strip()
    return None


def _extract_knowledge_area(lines: list[str]) -> str:
    tagged = _extract_tagged_block(
        lines,
        [
            "área do conhecimento",
            "area do conhecimento",
            "grande área",
            "grande area",
            "área:",
            "area:",
        ],
    )
    if tagged:
        return tagged

    for line in lines:
        lowered = line.lower()
        if any(
            token in lowered
            for token in [
                "engenharia",
                "educação",
                "educacao",
                "fisioterapia",
                "terapia ocupacional",
                "química",
                "quimica",
                "saúde",
                "saude",
                "computação",
                "computacao",
            ]
        ):
            return line

    return "Não identificado automaticamente a partir do currículo extraído."


def _extract_keywords(lines: list[str], summary: str, knowledge_area: str) -> str:
    tagged = _extract_tagged_block(
        lines, ["palavras-chave", "palavras chave", "temas de interesse"]
    )
    if tagged:
        return tagged

    candidates: list[str] = []
    for line in lines[:120]:
        if len(line) > 180:
            continue
        if line.count("/") >= 2:
            return line
        if re.search(
            r"\b(pesquisa|desenvolvimento|processos|saúde|educação|familia|energia)\b",
            line.lower(),
        ):
            candidates.append(line)
        if len(candidates) >= 3:
            break

    if candidates:
        return " | ".join(candidates)

    fallback = ", ".join(
        part.strip()
        for part in [knowledge_area, summary[:140].rsplit(" ", 1)[0]]
        if part.strip()
    )
    return fallback or "Definir manualmente com base no trecho-base."


def _build_summary_paragraph(lines: list[str], nome: str) -> str:
    narrative_lines = [
        line
        for line in lines
        if len(line) >= 60
        and not line.lower().startswith("campo ")
        and "http" not in line.lower()
    ]
    joined = " ".join(narrative_lines)
    joined = re.sub(r"\s+", " ", joined).strip()
    if not joined:
        return (
            f"{_normalize_display_name(nome)} possui currículo Lattes disponível e requer síntese editorial manual, "
            "pois a extração automática não encontrou um bloco narrativo consistente."
        )

    sentences = re.split(r"(?<=[.!?])\s+", joined)
    selected: list[str] = []
    current_length = 0
    for sentence in sentences:
        clean_sentence = sentence.strip()
        if len(clean_sentence) < 40:
            continue
        if current_length + len(clean_sentence) > 680 and selected:
            break
        selected.append(clean_sentence)
        current_length += len(clean_sentence) + 1
        if current_length >= 420:
            break

    summary = " ".join(selected).strip()
    if not summary:
        summary = joined[:680].rsplit(" ", 1)[0]
    return summary


def _infer_ods(summary: str, keywords: str, knowledge_area: str) -> list[str]:
    text = " ".join([summary, keywords, knowledge_area]).lower()
    inferred: list[str] = []

    mapping = [
        (
            "ODS 3 - Saúde e bem-estar",
            ["saúde", "saude", "terapia", "reabilita", "bem-estar", "hospital"],
        ),
        (
            "ODS 4 - Educação de qualidade",
            ["educação", "educacao", "ensino", "aprendiz", "escola", "universidade"],
        ),
        (
            "ODS 7 - Energia limpa e acessível",
            [
                "energia",
                "biocombust",
                "etanol",
                "transição energética",
                "transicao energetica",
            ],
        ),
        (
            "ODS 9 - Indústria, inovação e infraestrutura",
            [
                "processos",
                "engenharia",
                "inovação",
                "inovacao",
                "modelagem",
                "simulação",
                "simulacao",
            ],
        ),
        (
            "ODS 10 - Redução das desigualdades",
            [
                "vulnerabilidade",
                "desigual",
                "inclus",
                "assistiva",
                "autonomia",
                "família",
                "familia",
            ],
        ),
        (
            "ODS 13 - Ação contra a mudança global do clima",
            ["sustentabilidade", "clima", "emissões", "emissoes", "ambiental"],
        ),
    ]

    for label, tokens in mapping:
        if any(token in text for token in tokens):
            inferred.append(label)

    return inferred[:3]


def _build_source_excerpt(lines: list[str]) -> str:
    meaningful = [
        line
        for line in lines
        if not line.lower().startswith("campo ")
        and len(line) > 20
        and "http" not in line.lower()
    ]
    excerpt = " ".join(meaningful)
    excerpt = re.sub(r"\s+", " ", excerpt).strip()
    if not excerpt:
        return "Trecho-base indisponível na extração automática."
    return excerpt[:1800].rsplit(" ", 1)[0]


def build_export_summary_context(
    texto_extraido: str, nome: str, *, has_profile_photo: bool = False
) -> ExportSummaryContext:
    lines = _clean_text_lines(texto_extraido)
    summary = _build_summary_paragraph(lines, nome)
    knowledge_area = _extract_knowledge_area(lines)
    keywords = _extract_keywords(lines, summary, knowledge_area)
    social_links = _extract_urls(lines)
    focus_topics = (
        keywords if len(keywords) <= 500 else keywords[:500].rsplit(" ", 1)[0]
    )
    ods = _infer_ods(summary, keywords, knowledge_area)
    excerpt = _build_source_excerpt(lines)

    return ExportSummaryContext(
        summary_paragraph=summary,
        keywords=keywords,
        knowledge_area=knowledge_area,
        social_links=social_links,
        focus_topics=focus_topics,
        ods=ods,
        source_excerpt=excerpt,
        has_profile_photo=has_profile_photo,
    )


def _write_json_bytes(
    *,
    nome: str,
    ultima_atualizacao: date,
    cache_status: str | None,
    pdf_filename: str,
    pdf_storage_path: str,
    pdf_download_url: str,
    texto_extraido: str,
    context: ExportSummaryContext,
) -> bytes:
    payload = {
        "nome": nome,
        "ultima_atualizacao_curriculo": ultima_atualizacao.isoformat(),
        "cache_status": cache_status,
        "arquivo_pdf": pdf_filename,
        "storage_path": pdf_storage_path,
        "download_pdf_url": pdf_download_url,
        "template_docx": _SUMMARY_TEMPLATE_NAME,
        "gerado_em": datetime.now().isoformat(),
        "resumo_contextual": asdict(context),
        "foto_extraida": context.has_profile_photo,
        "texto_extraido": texto_extraido,
    }
    return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")


def _write_html_bytes(
    *,
    nome: str,
    ultima_atualizacao: date,
    cache_status: str | None,
    context: ExportSummaryContext,
    photo_bytes: bytes | None,
    photo_content_type: str | None,
) -> bytes:
    social_html = "".join(f"<li>{link}</li>" for link in context.social_links)
    if not social_html:
        social_html = "<li>Não identificado automaticamente.</li>"
    ods_html = "".join(f"<li>{item}</li>" for item in context.ods)
    if not ods_html:
        ods_html = "<li>Não inferido automaticamente.</li>"
    photo_html = ""
    if photo_bytes:
        content_type = photo_content_type or "image/png"
        photo_src = b64encode(photo_bytes).decode("ascii")
        photo_html = (
            f'<div class="portrait"><img alt="Foto de perfil de {escape(nome)}" '
            f'src="data:{content_type};base64,{photo_src}" /></div>'
        )

    html = f"""<!DOCTYPE html>
<html lang=\"pt-BR\">
  <head>
    <meta charset=\"utf-8\" />
    <meta name=\"viewport\" content=\"width=device-width, initial-scale=1\" />
    <title>Perfil Vitrine - {escape(nome)}</title>
    <style>
      :root {{
        color-scheme: light;
        --bg: #f7f1e8;
        --panel: #fffdf9;
        --line: #decfb6;
        --ink: #1f2a37;
        --accent: #9a3412;
      }}
      body {{ margin: 0; font-family: Georgia, \"Times New Roman\", serif; background: linear-gradient(180deg, #fbf7ef 0%, var(--bg) 100%); color: var(--ink); }}
      main {{ width: min(980px, calc(100% - 32px)); margin: 32px auto; padding: 32px; background: var(--panel); border: 1px solid var(--line); border-radius: 28px; box-shadow: 0 24px 80px -56px rgba(15, 23, 42, 0.55); }}
      h1, h2 {{ margin-top: 0; }}
            .hero {{ display: grid; grid-template-columns: minmax(180px, 240px) 1fr; gap: 24px; align-items: start; margin-bottom: 24px; }}
      .tag {{ display: inline-block; margin-bottom: 12px; padding: 6px 12px; background: rgba(154, 52, 18, 0.12); color: var(--accent); border-radius: 999px; font-weight: 700; font-size: 0.85rem; }}
      .grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(260px, 1fr)); gap: 16px; }}
      .card {{ padding: 18px; border-radius: 18px; border: 1px solid var(--line); background: rgba(255,255,255,0.76); }}
            .portrait {{ border-radius: 22px; overflow: hidden; border: 1px solid var(--line); background: linear-gradient(180deg, #f7ede0 0%, #fff 100%); min-height: 220px; display: flex; align-items: center; justify-content: center; }}
            .portrait img {{ display: block; width: 100%; height: auto; object-fit: cover; }}
            .muted {{ color: #6b7280; }}
      p, li {{ line-height: 1.7; }}
      pre {{ white-space: pre-wrap; font-family: inherit; margin: 0; }}
    </style>
  </head>
  <body>
    <main>
      <span class=\"tag\">Perfil Vitrine resumido</span>
            <div class=\"hero\">
                {photo_html or '<div class="portrait"><p class="muted">Foto não extraída automaticamente.</p></div>'}
                <div>
                    <h1>{escape(nome)}</h1>
                    <p><strong>Área do conhecimento:</strong> {escape(context.knowledge_area)}</p>
                    <p class=\"muted\">Última atualização do currículo: {ultima_atualizacao.isoformat()} | Origem: {escape(cache_status or 'não informada')}</p>
                    <div class=\"card\"><h2>Parágrafo síntese</h2><p>{escape(context.summary_paragraph)}</p></div>
                </div>
            </div>
      <div class=\"grid\">
                <section class=\"card\"><h2>Palavras-chave</h2><p>{escape(context.keywords)}</p></section>
        <section class=\"card\"><h2>ODS</h2><ul>{ods_html}</ul></section>
        <section class=\"card\"><h2>Redes e plataformas</h2><ul>{social_html}</ul></section>
                <section class=\"card\"><h2>Temas de interesse</h2><p>{escape(context.focus_topics)}</p></section>
                <section class=\"card\"><h2>Trecho-base</h2><pre>{escape(context.source_excerpt)}</pre></section>
      </div>
    </main>
  </body>
</html>
"""
    return html.encode("utf-8")


def _write_csv_bytes(
    *,
    nome: str,
    ultima_atualizacao: date,
    cache_status: str | None,
    pdf_filename: str,
    pdf_storage_path: str,
    pdf_download_url: str,
    context: ExportSummaryContext,
) -> bytes:
    buffer = StringIO()
    writer = csv.DictWriter(
        buffer,
        fieldnames=[
            "nome",
            "ultima_atualizacao_curriculo",
            "cache_status",
            "arquivo_pdf",
            "storage_path",
            "download_pdf_url",
            "area_do_conhecimento",
            "palavras_chave",
            "paragrafo_sintese",
            "ods",
            "redes_sociais",
        ],
    )
    writer.writeheader()
    writer.writerow(
        {
            "nome": nome,
            "ultima_atualizacao_curriculo": ultima_atualizacao.isoformat(),
            "cache_status": cache_status or "",
            "arquivo_pdf": pdf_filename,
            "storage_path": pdf_storage_path,
            "download_pdf_url": pdf_download_url,
            "area_do_conhecimento": context.knowledge_area,
            "palavras_chave": context.keywords,
            "paragrafo_sintese": context.summary_paragraph,
            "ods": " | ".join(context.ods),
            "redes_sociais": " | ".join(context.social_links),
        }
    )
    return buffer.getvalue().encode("utf-8")


def _write_docx_bytes(
    *,
    nome: str,
    ultima_atualizacao: date,
    cache_status: str | None,
    context: ExportSummaryContext,
    photo_bytes: bytes | None,
) -> bytes:
    try:
        from docx import Document
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt, RGBColor
    except ImportError as exc:
        raise ValueError(
            "Dependência 'python-docx' não encontrada. Atualize o backend com 'uv sync'."
        ) from exc

    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Aptos"
    normal_style.font.size = Pt(10.5)

    accent = RGBColor(0x9A, 0x34, 0x12)
    muted = RGBColor(0x5B, 0x63, 0x70)

    def add_heading(text: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = accent

    def add_body(text: str) -> None:
        paragraph = document.add_paragraph(text)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.25

    def add_list(items: list[str], empty_message: str) -> None:
        if not items:
            add_body(empty_message)
            return
        for item in items:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.add_run(item)

    overline = document.add_paragraph()
    overline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    overline_run = overline.add_run("PERFIL VITRINE")
    overline_run.bold = True
    overline_run.font.size = Pt(10)
    overline_run.font.color.rgb = accent

    hero = document.add_table(rows=1, cols=2)
    hero.autofit = False
    hero.columns[0].width = Cm(4.8)
    hero.columns[1].width = Cm(11.8)

    photo_cell = hero.rows[0].cells[0]
    photo_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    photo_paragraph = photo_cell.paragraphs[0]
    photo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if photo_bytes:
        photo_paragraph.add_run().add_picture(BytesIO(photo_bytes), width=Cm(4.2))
    else:
        placeholder = photo_paragraph.add_run("Foto não extraída automaticamente")
        placeholder.italic = True
        placeholder.font.color.rgb = muted

    info_cell = hero.rows[0].cells[1]
    info_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    title = info_cell.paragraphs[0]
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run(nome)
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x37)

    subtitle = info_cell.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(4)
    subtitle_run = subtitle.add_run(context.knowledge_area)
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = accent

    meta = info_cell.add_paragraph()
    meta.paragraph_format.space_after = Pt(8)
    meta_run = meta.add_run(
        f"Última atualização do currículo: {ultima_atualizacao.isoformat()} | Origem: {cache_status or 'não informada'}"
    )
    meta_run.italic = True
    meta_run.font.size = Pt(9.5)
    meta_run.font.color.rgb = muted

    synthesis_label = info_cell.add_paragraph()
    synthesis_label.paragraph_format.space_after = Pt(2)
    synthesis_label_run = synthesis_label.add_run("Parágrafo síntese")
    synthesis_label_run.bold = True
    synthesis_label_run.font.color.rgb = accent
    synthesis = info_cell.add_paragraph(context.summary_paragraph)
    synthesis.paragraph_format.line_spacing = 1.25
    synthesis.paragraph_format.space_after = Pt(0)

    add_heading("Palavras-chave")
    add_body(context.keywords)

    add_heading("ODS")
    add_list(
        context.ods,
        "Não inferido automaticamente a partir do currículo extraído.",
    )

    add_heading("Redes sociais e plataformas educacionais")
    add_list(
        context.social_links,
        "Não identificadas automaticamente no currículo extraído.",
    )

    add_heading("Temas de interesse")
    add_body(context.focus_topics)

    add_heading("Trecho-base para edição")
    add_body(context.source_excerpt)

    add_heading("Metadados")
    add_body(
        f"Modelo editorial: {_SUMMARY_TEMPLATE_NAME}. Foto extraída automaticamente: {'sim' if context.has_profile_photo else 'não'}."
    )

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _merge_source_texts(pdf_text: str, html_text: str | None) -> str:
    if not html_text:
        return pdf_text
    if not pdf_text:
        return html_text

    merged_lines: list[str] = []
    seen: set[str] = set()
    for source in [pdf_text, html_text]:
        for line in _clean_text_lines(source):
            key = line.lower()
            if key in seen:
                continue
            seen.add(key)
            merged_lines.append(line)
    return "\n".join(merged_lines)


def _build_artifact_contents(
    *,
    nome: str,
    ultima_atualizacao: date,
    pdf_filename: str,
    pdf_storage_path: str,
    pdf_download_url: str,
    pdf_bytes: bytes,
    requested_formats: list[OutputFormat],
    cache_status: str | None,
    html_text: str | None = None,
    photo_bytes: bytes | None = None,
    photo_content_type: str | None = None,
) -> tuple[dict[str, tuple[str, bytes, str]], ExportSummaryContext, int]:
    texto_pdf = _extrair_texto_pdf_bytes(pdf_bytes).strip()
    texto_extraido = _merge_source_texts(texto_pdf, html_text).strip()
    context = build_export_summary_context(
        texto_extraido,
        nome,
        has_profile_photo=photo_bytes is not None,
    )
    specs = _artifact_file_specs(nome, ultima_atualizacao)
    contents: dict[str, tuple[str, bytes, str]] = {}

    for file_format in requested_formats:
        filename, content_type = specs[file_format]
        if file_format == "pdf":
            contents[file_format] = (filename, pdf_bytes, content_type)
        elif file_format == "json":
            contents[file_format] = (
                filename,
                _write_json_bytes(
                    nome=nome,
                    ultima_atualizacao=ultima_atualizacao,
                    cache_status=cache_status,
                    pdf_filename=pdf_filename,
                    pdf_storage_path=pdf_storage_path,
                    pdf_download_url=pdf_download_url,
                    texto_extraido=texto_extraido,
                    context=context,
                ),
                content_type,
            )
        elif file_format == "html":
            contents[file_format] = (
                filename,
                _write_html_bytes(
                    nome=nome,
                    ultima_atualizacao=ultima_atualizacao,
                    cache_status=cache_status,
                    context=context,
                    photo_bytes=photo_bytes,
                    photo_content_type=photo_content_type,
                ),
                content_type,
            )
        elif file_format == "csv":
            contents[file_format] = (
                filename,
                _write_csv_bytes(
                    nome=nome,
                    ultima_atualizacao=ultima_atualizacao,
                    cache_status=cache_status,
                    pdf_filename=pdf_filename,
                    pdf_storage_path=pdf_storage_path,
                    pdf_download_url=pdf_download_url,
                    context=context,
                ),
                content_type,
            )
        elif file_format == "docx":
            contents[file_format] = (
                filename,
                _write_docx_bytes(
                    nome=nome,
                    ultima_atualizacao=ultima_atualizacao,
                    cache_status=cache_status,
                    context=context,
                    photo_bytes=photo_bytes,
                ),
                content_type,
            )

    return contents, context, len(texto_extraido)


def _artifact_lookup(folder: str) -> dict[str, object]:
    return {item.filename: item for item in list_storage_files(folder)}


def _load_manifest(folder: str) -> dict[str, object] | None:
    object_path = f"{folder}/{_MANIFEST_FILENAME}"
    data = download_storage_file_bytes(object_path)
    if not data:
        return None
    try:
        return json.loads(data.decode("utf-8"))
    except (UnicodeDecodeError, json.JSONDecodeError):
        return None


def _coerce_int(value: object, default: int = 0) -> int:
    if isinstance(value, bool):
        return int(value)
    if isinstance(value, int):
        return value
    if isinstance(value, float):
        return int(value)
    if isinstance(value, str):
        try:
            return int(value.strip())
        except ValueError:
            return default
    try:
        return int(str(value))
    except (TypeError, ValueError):
        return default


def _artifact_format_by_filename(filename: str) -> str | None:
    if filename.startswith("perfil-vitrine-") and filename.endswith(".docx"):
        return "docx"
    if filename == "curriculo-lattes.pdf":
        return "pdf"
    if filename == "dados-extraidos.json":
        return "json"
    if filename == "curriculo-lattes.html":
        return "html"
    if filename == "dados-extraidos.csv":
        return "csv"
    if filename.endswith(".zip"):
        return "zip"
    return None


def _build_bundle_from_storage(
    *,
    nome: str,
    ultima_atualizacao: date,
    folder: str,
    output_label: str,
    output_format: OutputFormat,
    requested_formats: list[OutputFormat],
    extracted_text_length: int,
    template_name: str | None,
    artifacts_cache_status: str,
) -> GeneratedArtifactBundle:
    files = list_storage_files(folder)
    files_by_name = {item.filename: item for item in files}
    specs = _artifact_file_specs(nome, ultima_atualizacao)

    generated_files: list[GeneratedArtifact] = []
    for file_format in requested_formats:
        filename, content_type = specs[file_format]
        storage_file = files_by_name[filename]
        generated_files.append(
            GeneratedArtifact(
                format=file_format,
                filename=filename,
                relative_path=storage_file.object_path,
                download_url=storage_file.download_url,
                content_type=content_type,
            )
        )

    zip_file = None
    for item in files:
        if item.filename.endswith(".zip"):
            zip_file = GeneratedArtifact(
                format="zip",
                filename=item.filename,
                relative_path=item.object_path,
                download_url=item.download_url,
                content_type=item.content_type or "application/zip",
            )
            break

    return GeneratedArtifactBundle(
        output_format=output_format,
        output_directory=folder,
        output_label=output_label,
        generated_files=generated_files,
        extracted_text_length=extracted_text_length,
        template_name=template_name,
        zip_file=zip_file,
        artifacts_cache_status=artifacts_cache_status,
    )


def _build_manifest_payload(
    *,
    nome: str,
    ultima_atualizacao: date,
    output_directory: str,
    output_label: str,
    extracted_text_length: int,
    generated_files: list[GeneratedArtifact],
    has_profile_photo: bool,
) -> bytes:
    payload = {
        "nome": nome,
        "ultima_atualizacao_curriculo": ultima_atualizacao.isoformat(),
        "output_directory": output_directory,
        "output_label": output_label,
        "template_name": _SUMMARY_TEMPLATE_NAME,
        "template_version": _artifact_template_version(),
        "generated_at": datetime.now().isoformat(),
        "extracted_text_length": extracted_text_length,
        "has_profile_photo": has_profile_photo,
        "generated_files": [asdict(item) for item in generated_files],
    }
    return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")


def _ensure_package_zip(
    *,
    folder: str,
    nome: str,
    ultima_atualizacao: date,
    force_regenerate: bool = False,
) -> GeneratedArtifact | None:
    zip_filename = _package_zip_filename(nome, ultima_atualizacao)
    existing_files = list_storage_files(folder)
    existing_by_name = {item.filename: item for item in existing_files}

    if not force_regenerate and zip_filename in existing_by_name:
        item = existing_by_name[zip_filename]
        return GeneratedArtifact(
            format="zip",
            filename=item.filename,
            relative_path=item.object_path,
            download_url=item.download_url,
            content_type=item.content_type or "application/zip",
        )

    zip_buffer = BytesIO()
    with zipfile.ZipFile(
        zip_buffer, mode="w", compression=zipfile.ZIP_DEFLATED
    ) as zip_file:
        for item in existing_files:
            if item.filename in {_MANIFEST_FILENAME, zip_filename}:
                continue
            file_bytes = download_storage_file_bytes(item.object_path)
            if file_bytes is None:
                continue
            zip_file.writestr(item.filename, file_bytes)

    upload_result = upload_file_bytes(
        zip_filename,
        zip_buffer.getvalue(),
        content_type="application/zip",
        folder=folder,
    )
    return GeneratedArtifact(
        format="zip",
        filename=zip_filename,
        relative_path=upload_result.object_path,
        download_url=upload_result.download_url,
        content_type="application/zip",
    )


def ensure_curriculo_artifacts(
    *,
    nome: str,
    ultima_atualizacao: date,
    pdf_filename: str,
    pdf_storage_path: str,
    pdf_download_url: str,
    pdf_bytes: bytes,
    output_format: OutputFormat,
    cache_status: str | None = None,
    html_text: str | None = None,
    photo_bytes: bytes | None = None,
    photo_content_type: str | None = None,
) -> GeneratedArtifactBundle:
    folder = build_curriculo_storage_folder(nome, ultima_atualizacao)
    output_label = build_curriculo_output_label(nome, ultima_atualizacao)
    requested_formats = expand_output_formats(output_format)
    specs = _artifact_file_specs(nome, ultima_atualizacao)
    manifest = _load_manifest(folder) or {}
    existing = _artifact_lookup(folder)

    missing_formats: list[OutputFormat] = [
        file_format
        for file_format in requested_formats
        if specs[file_format][0] not in existing
    ]

    extracted_text_length = _coerce_int(manifest.get("extracted_text_length"), 0)

    if missing_formats:
        contents, _, extracted_text_length = _build_artifact_contents(
            nome=nome,
            ultima_atualizacao=ultima_atualizacao,
            pdf_filename=pdf_filename,
            pdf_storage_path=pdf_storage_path,
            pdf_download_url=pdf_download_url,
            pdf_bytes=pdf_bytes,
            requested_formats=missing_formats,
            cache_status=cache_status,
            html_text=html_text,
            photo_bytes=photo_bytes,
            photo_content_type=photo_content_type,
        )
        for _, (filename, file_bytes, content_type) in contents.items():
            upload_file_bytes(
                filename,
                file_bytes,
                content_type=content_type,
                folder=folder,
            )

        refreshed_files = list_storage_files(folder)
        generated_files = [
            GeneratedArtifact(
                format=_artifact_format_by_filename(item.filename) or "unknown",
                filename=item.filename,
                relative_path=item.object_path,
                download_url=item.download_url,
                content_type=item.content_type or "application/octet-stream",
            )
            for item in refreshed_files
            if item.filename != _MANIFEST_FILENAME
            and _artifact_format_by_filename(item.filename) not in {None, "zip"}
        ]
        manifest_bytes = _build_manifest_payload(
            nome=nome,
            ultima_atualizacao=ultima_atualizacao,
            output_directory=folder,
            output_label=output_label,
            extracted_text_length=extracted_text_length,
            generated_files=generated_files,
            has_profile_photo=photo_bytes is not None,
        )
        upload_file_bytes(
            _MANIFEST_FILENAME,
            manifest_bytes,
            content_type="application/json",
            folder=folder,
        )
        zip_file = _ensure_package_zip(
            folder=folder,
            nome=nome,
            ultima_atualizacao=ultima_atualizacao,
            force_regenerate=True,
        )
        bundle = _build_bundle_from_storage(
            nome=nome,
            ultima_atualizacao=ultima_atualizacao,
            folder=folder,
            output_label=output_label,
            output_format=output_format,
            requested_formats=requested_formats,
            extracted_text_length=extracted_text_length,
            template_name=_SUMMARY_TEMPLATE_NAME,
            artifacts_cache_status="miss",
        )
        return GeneratedArtifactBundle(
            output_format=bundle.output_format,
            output_directory=bundle.output_directory,
            output_label=bundle.output_label,
            generated_files=bundle.generated_files,
            extracted_text_length=bundle.extracted_text_length,
            template_name=bundle.template_name,
            zip_file=zip_file,
            artifacts_cache_status="miss",
        )

    zip_file = _ensure_package_zip(
        folder=folder,
        nome=nome,
        ultima_atualizacao=ultima_atualizacao,
        force_regenerate=False,
    )
    bundle = _build_bundle_from_storage(
        nome=nome,
        ultima_atualizacao=ultima_atualizacao,
        folder=folder,
        output_label=str(manifest.get("output_label") or output_label),
        output_format=output_format,
        requested_formats=requested_formats,
        extracted_text_length=extracted_text_length,
        template_name=str(manifest.get("template_name") or _SUMMARY_TEMPLATE_NAME),
        artifacts_cache_status="hit",
    )
    return GeneratedArtifactBundle(
        output_format=bundle.output_format,
        output_directory=bundle.output_directory,
        output_label=bundle.output_label,
        generated_files=bundle.generated_files,
        extracted_text_length=bundle.extracted_text_length,
        template_name=bundle.template_name,
        zip_file=zip_file or bundle.zip_file,
        artifacts_cache_status="hit",
    )


def upload_batch_zip(
    *,
    batch_folder: str,
    batch_filename: str,
    entries: list[tuple[str, bytes]],
) -> GeneratedArtifact:
    zip_buffer = BytesIO()
    with zipfile.ZipFile(
        zip_buffer, mode="w", compression=zipfile.ZIP_DEFLATED
    ) as zip_file:
        for arcname, file_bytes in entries:
            zip_file.writestr(arcname, file_bytes)

    upload_result = upload_file_bytes(
        batch_filename,
        zip_buffer.getvalue(),
        content_type="application/zip",
        folder=batch_folder,
    )
    return GeneratedArtifact(
        format="zip",
        filename=batch_filename,
        relative_path=upload_result.object_path,
        download_url=upload_result.download_url,
        content_type="application/zip",
    )


def download_artifact_bytes(artifact: GeneratedArtifact) -> bytes | None:
    return download_storage_file_bytes(artifact.relative_path)


def artifacts_to_payload(bundle: GeneratedArtifactBundle) -> dict[str, object]:
    return {
        "output_format": bundle.output_format,
        "output_directory": bundle.output_directory,
        "output_label": bundle.output_label,
        "generated_files": [asdict(file) for file in bundle.generated_files],
        "extracted_text_length": bundle.extracted_text_length,
        "template_name": bundle.template_name,
        "artifacts_cache_status": bundle.artifacts_cache_status,
        "zip_arquivo": bundle.zip_file.filename if bundle.zip_file else None,
        "zip_storage_path": bundle.zip_file.relative_path if bundle.zip_file else None,
        "zip_download_url": bundle.zip_file.download_url if bundle.zip_file else None,
    }
